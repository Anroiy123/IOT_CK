from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
SOURCE_MD = REPORTS / "BaoCao_IOT_Nhom12.md"
OUT_DOCX = ROOT / "bao_cao_gesture_recognition_iot_ck.docx"

BODY_FONT = "Times New Roman"


def set_run_font(run, size=11, bold=None, italic=None):
    run.font.name = BODY_FONT
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def clean_inline(text: str) -> str:
    text = text.replace("\\[", "[").replace("\\]", "]").replace("\\*", "*")
    text = text.replace("**", "").replace("*", "")
    return text.strip()


def set_para_format(p, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, after=6, line=1.5):
    p.alignment = align
    p.paragraph_format.line_spacing = line
    p.paragraph_format.space_after = Pt(after)
    if first_line:
        p.paragraph_format.first_line_indent = Cm(0.8)


def add_p(doc, text="", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True, size=11, bold=False, italic=False, after=6):
    p = doc.add_paragraph()
    set_para_format(p, align=align, first_line=first_line, after=after)
    if text:
        r = p.add_run(clean_inline(text))
        set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level):
    text = clean_inline(text)
    p = doc.add_paragraph()
    if level == 1:
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, after=10, line=1.2)
        r = p.add_run(text.upper())
        set_run_font(r, size=15, bold=True)
    elif level == 2:
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, after=6, line=1.25)
        r = p.add_run(text)
        set_run_font(r, size=13, bold=True)
    else:
        set_para_format(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False, after=4, line=1.25)
        r = p.add_run(text)
        set_run_font(r, size=12, bold=True)
    return p


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Trang ")
    set_run_font(r, size=10)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def setup_document(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)
    add_page_number(sec)
    for style_name in ["Normal", "List Bullet", "List Number"]:
        st = doc.styles[style_name]
        st.font.name = BODY_FONT
        st._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        st.font.size = Pt(11)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_after = Pt(6)
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_cell(cell, text, *, bold=False, center=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(clean_inline(str(text)))
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    borders = table._tbl.tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table._tbl.tblPr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = "w:" + edge
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "4")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "7F7F7F")


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)
    for i, val in enumerate(rows[0]):
        set_cell(table.rows[0].cells[i], val, bold=True, center=True, size=10)
        shade_cell(table.rows[0].cells[i])
    repeat_table_header(table.rows[0])
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(len(rows[0])):
            set_cell(cells[i], row[i] if i < len(row) else "", size=10 if len(rows[0]) >= 5 else 11)
    add_p(doc, "", first_line=False, after=2)


def add_table_caption(doc, text):
    add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=11, bold=True, after=3)


def add_figure_caption(doc, text):
    add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=11, italic=True, after=8)


def add_figure(doc, src, caption):
    num = re.search(r"Hình\s+\d+(\.\d+)?", caption)
    fig_name = num.group(0) if num else "Hình"
    add_p(doc, f"{fig_name} được chèn dưới đây để minh họa trực quan cho phần mô tả và phân tích ngay trong mục này.", first_line=True)
    path = (ROOT / src).resolve()
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(path), width=Inches(5.9))
    add_figure_caption(doc, caption)


def cover(doc):
    for text in [
        "BỘ KHOA HỌC VÀ CÔNG NGHỆ",
        "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG",
        "CƠ SỞ TẠI THÀNH PHỐ HỒ CHÍ MINH",
        "----------o0o----------",
    ]:
        add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=13, bold=True)
    for _ in range(4):
        add_p(doc, "", first_line=False, after=10)
    add_p(doc, "BÁO CÁO", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=24, bold=True)
    add_p(doc, "ĐỒ ÁN MÔN HỌC", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False, size=18, bold=True)
    add_p(
        doc,
        "THIẾT KẾ THIẾT BỊ ĐIỀU KHIỂN BẰNG CỬ CHỈ TAY SỬ DỤNG DEEP LEARNING CHO ĐIỀU KHIỂN THÔNG MINH XE TỰ HÀNH",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
        size=15,
        bold=True,
    )
    add_p(doc, "", first_line=False)
    add_table(doc, [
        ["Thông tin", "Nội dung"],
        ["Môn học", "Phát triển ứng dụng IOT"],
        ["Giảng viên hướng dẫn", "Đàm Minh Lịnh"],
        ["Nhóm thực hiện", "Nhóm 12"],
        ["Lớp", "D22CQPTUD01-N"],
    ])
    add_table(doc, [
        ["MSSV", "Họ tên", "Email", "Vai trò"],
        ["N22DCPT035", "Trần Quang Hùng", "n22dcpt035@student.ptithcm.edu.vn", "Nhóm trưởng"],
        ["N22DCPT052", "Vũ Quang Long", "n22dcpt052@student.ptithcm.edu.vn", "Thành viên"],
    ])
    add_p(doc, "TP. Hồ Chí Minh, 15 / 6 / 2026", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    doc.add_page_break()


def transform_markdown(md: str) -> str:
    md = md.replace("\r\n", "\n")
    start = md.index("# **LỜI CẢM ƠN**")
    refs_start = md.index("# **TÀI LIỆU THAM KHẢO**")
    app_a_start = md.index("# **PHỤ LỤC A.")
    before_refs = md[start:refs_start]
    refs = md[refs_start:app_a_start]

    ch3 = before_refs.index("# **CHƯƠNG 3.")
    source_check = """
## **2.11 Kiểm tra nguồn khoa học và quản lý trích dẫn**

Theo lưu ý của giảng viên, nhóm ưu tiên các bài báo có DOI, xuất bản bởi IEEE,
ACM, Elsevier, Springer hoặc tạp chí/hội nghị có trang tra cứu chỉ mục rõ ràng.
Danh mục tài liệu tham khảo cuối báo cáo gồm 29 nguồn, trong đó các nguồn [1]-[4]
được dùng cho nền tảng HGR/HCI/HRI, các nguồn [5]-[14] dùng cho MediaPipe,
MobileNet, CNN-LSTM, 3D-CNN và dataset cử chỉ, các nguồn [15]-[24] dùng cho
ứng dụng robot/control, và các nguồn [25]-[29] dùng cho tiêu chí đánh giá,
cloud API, FastAPI, WebSocket và chuẩn truyền thông.

Quy trình quản lý tài liệu được đề xuất là nhập DOI hoặc URL chính thức vào
Zotero, kiểm tra metadata gồm tác giả, năm, tiêu đề, venue, DOI và URL, sau đó
đối chiếu nhanh venue trên SCImago, Scopus Sources hoặc Web of Science Journal
Info nếu đó là tạp chí. Với conference paper, nhóm ưu tiên trang chính thức của
IEEE Xplore, ACM Digital Library, CVF Open Access hoặc DOI resolver. ResearchGate
chỉ nên dùng để tìm bản full text hoặc metadata phụ, không dùng làm bằng chứng
chỉ mục.

Việc đưa mục kiểm tra nguồn vào Chương 2 giúp báo cáo không biến phụ lục thành
phần dài quá mức. Các nguồn được trích trực tiếp trong phần thân bằng định dạng
IEEE numeric [n]; danh mục cuối báo cáo cung cấp DOI/URL để người đọc nhập vào
Zotero hoặc Mendeley khi cần.

| **Nhóm nguồn** | **Nguồn tiêu biểu** | **Vai trò trong báo cáo** | **Kênh kiểm tra ưu tiên** |
|:---|:---|:---|:---|
| Tổng quan HGR/HRI | [1], [2], [3], [4] | Xây dựng bối cảnh nghiên cứu, khó khăn của nhận dạng cử chỉ và yêu cầu hệ thống robot | DOI, trang tạp chí, SCImago/WoS/Scopus nếu là journal |
| Landmark và tiền xử lý | [5] | Biện minh lựa chọn MediaPipe Hands để phát hiện bàn tay và crop vùng quan tâm | arXiv chính thức, repository/tài liệu MediaPipe |
| Mô hình nhẹ | [6], [7] | Biện minh MobileNetV3Small và họ MobileNet cho bài toán real-time | CVF/IEEE DOI |
| Dataset và cử chỉ động | [8], [9], [10], [12], [13], [14] | So sánh dữ liệu nội bộ với dataset/cách học temporal trong nghiên cứu trước | IEEE/CVF, Elsevier DOI |
| Robot/control | [15], [16], [17], [18], [19], [20], [21] | Liên hệ đầu ra HGR với điều khiển UAV, robot arm, mobile robot và collaborative assembly | IEEE Xplore, ACM DL, CVF |
| Cảm biến và phương pháp thay thế | [22], [23], [24] | Làm rõ phạm vi đề tài chỉ dùng camera RGB, không dùng radar hoặc Leap Motion | DOI, trang tạp chí |
| Đánh giá và triển khai | [25], [26], [27], [28], [29] | Chuẩn hóa metric, cloud deployment, FastAPI và WebSocket | DOI/RFC/trang tài liệu chính thức |

*Bảng 2.3. Ma trận nguồn khoa học và vai trò sử dụng trong báo cáo*

Bảng 2.3 cho thấy danh mục tài liệu không chỉ được liệt kê ở cuối báo cáo mà
được tổ chức theo vai trò cụ thể. Nhóm tổng quan giúp xác định bài toán và
khoảng trống; nhóm mô hình giúp chọn baseline; nhóm robot/control giúp thiết kế
tầng lệnh; nhóm đánh giá giúp chọn metric và cách trình bày kết quả. Khi nộp
bản cuối, nhóm nên nhập toàn bộ DOI/URL vào Zotero, xuất bibliography theo IEEE
và rà lại số thứ tự sau khi chỉnh sửa nội dung.
"""
    ch3_extra = """
## **3.13 Chi tiết triển khai prototype phần cứng**

Prototype phần cứng được thiết kế theo hướng phân tách rõ khối xử lý thị giác
và khối chấp hành. Laptop hoặc gateway chịu trách nhiệm xử lý webcam, phát hiện
bàn tay, gọi cloud API và ra quyết định điều khiển. ESP32 chỉ nhận bản tin điều
khiển đã được gateway xác nhận, sau đó phát tín hiệu tới L298N cho hai động cơ
DC và PCA9685 cho các servo của tay robot. Cách phân tách này giúp firmware
nhẹ hơn, giảm rủi ro treo vi điều khiển do xử lý ảnh và phù hợp với năng lực
tính toán của ESP32.

Ở tầng nguồn, robot cần tách tải motor/servo khỏi nguồn logic. Motor DC tạo
dòng khởi động lớn, servo có thể tạo xung dòng khi đổi góc nhanh, trong khi
ESP32 và PCA9685 cần mức 5V/3.3V ổn định. Vì vậy thiết kế yêu cầu GND chung
giữa nguồn motor, nguồn servo và ESP32, nhưng đường cấp tải phải đủ dòng và
không lấy trực tiếp từ chân 5V yếu của board. Nếu GND không chung, tín hiệu PWM
hoặc mức logic điều khiển L298N có thể bị tham chiếu sai, làm motor phản hồi
không ổn định.

Hình 3.5 minh họa cụm tay robot trong giai đoạn lắp ráp và kiểm thử servo. Ảnh
này được đưa vào để liên hệ trực tiếp giữa mô hình điều khiển trong báo cáo và
cơ cấu chấp hành vật lý của đề tài.

<img src="lab Robot/z6995524077751_e0b45e7c2c2e79be95383dda926fa80d.jpg"
     alt="Cụm tay robot servo trong quá trình lắp ráp"
     width="720" />

*Hình 3.5. Cụm tay robot servo trong quá trình lắp ráp*

| **Rủi ro phần cứng** | **Biểu hiện** | **Biện pháp xử lý trong đề tài** |
|:---|:---|:---|
| Sụt áp khi motor chạy | ESP32 reset, WebSocket mất kết nối | Tách nguồn tải, dùng buck ổn định, kiểm tra GND chung |
| Servo đổi góc quá nhanh | Tay robot giật hoặc kẹt cơ khí | Giới hạn góc, cooldown servo, chỉ phát lệnh khi nhãn ổn định |
| Nhiễu đường PWM | Servo rung nhẹ hoặc sai vị trí | Dùng PCA9685, dây ngắn, cấp nguồn servo đủ dòng |
| Lệnh điều khiển lặp | Xe tiếp tục chạy khi mất camera/cloud | Deadman timeout và lệnh stop mặc định |
| Boot không ổn định | ESP32 không vào chương trình chính | Tránh kéo sai trạng thái các strapping pin, kiểm tra GPIO trước khi nạp |

*Bảng 3.3. Rủi ro phần cứng và biện pháp giảm thiểu*

## **3.14 Schema bản tin và nguyên tắc giao tiếp**

Gateway không gửi trực tiếp nhãn mô hình sang ESP32 theo dạng chuỗi tự do. Mỗi
bản tin điều khiển nên có cấu trúc gồm seq, gesture, command, confidence,
ttl_ms, timestamp_ms và token. Trường seq giúp firmware phát hiện bản tin cũ
hoặc lặp; confidence giúp log lại vì sao gateway phát hoặc chặn lệnh; ttl_ms
giới hạn thời gian bản tin còn hợp lệ; token tránh việc thiết bị lạ trong cùng
mạng gửi lệnh điều khiển tới ESP32. Khi ESP32 xử lý xong, bản tin ACK cần trả
về seq, status và latency_ms để gateway ghi log.

Thiết kế giao thức này phù hợp với đặc thù IoT vì đường truyền Wi-Fi trong lớp
học hoặc phòng thí nghiệm có thể không ổn định. Nếu gateway chỉ gửi lệnh dạng
forward hoặc left, hệ thống khó phân biệt lệnh mới với lệnh cũ. Nếu không có
TTL, một bản tin trễ vẫn có thể được thực thi sau khi người dùng đã đổi cử chỉ.
Nếu không có ACK, nhóm không đo được đoạn latency giữa gateway và ESP32. Vì vậy
schema bản tin là một phần của mô hình đề xuất chứ không chỉ là chi tiết lập
trình.

| **Trường** | **Kiểu dữ liệu** | **Ý nghĩa** | **Ví dụ** |
|:---|:---|:---|:---|
| seq | integer | Số thứ tự bản tin, tăng dần theo phiên | 1024 |
| gesture | string | Nhãn cử chỉ sau khi ổn định | stop |
| command | string | Lệnh điều khiển gửi tới firmware | motor_stop |
| confidence | float | Độ tin cậy của model sau softmax | 0.91 |
| ttl_ms | integer | Thời gian bản tin còn hợp lệ | 300 |
| timestamp_ms | integer | Thời điểm gateway tạo lệnh | 1710000000 |
| token | string | Chuỗi xác thực đơn giản giữa gateway và ESP32 | configured_secret |

*Bảng 3.4. Schema bản tin điều khiển gateway - ESP32*

## **3.15 Logic an toàn trước khi phát lệnh**

Safety filter trong đề tài gồm ba lớp. Lớp thứ nhất là ngưỡng confidence: nếu
xác suất cao nhất của softmax nhỏ hơn ngưỡng cấu hình, gateway không phát lệnh
điều khiển mà giữ trạng thái hiện tại hoặc gửi stop. Lớp thứ hai là voting theo
cửa sổ thời gian ngắn: một cử chỉ chỉ được chấp nhận khi xuất hiện ổn định trong
nhiều frame liên tiếp. Lớp thứ ba là deadman timeout: nếu gateway không nhận
được frame mới, cloud không phản hồi hoặc WebSocket lỗi quá thời gian cho phép,
firmware tự chuyển về trạng thái dừng.

Nguyên tắc này đặc biệt quan trọng với các lớp dễ nhầm như one, two, like và
dislike. Trong confusion matrix hiện có, các lớp này có độ chồng lấn cao hơn so
với no_gesture hoặc peace. Nếu mọi frame dự đoán đều trở thành lệnh, xe có thể
đổi hướng liên tục và tay robot có thể dao động. Safety filter hy sinh một phần
độ nhạy để đổi lấy hành vi ổn định hơn, phù hợp với mục tiêu demo IoT có phần
cứng thật.

## **3.16 Công thức đánh giá và thuật toán ổn định quyết định**

Đầu ra của CNN là vector logit z gồm tám phần tử, tương ứng tám lớp cử chỉ. Lớp
softmax chuyển logit thành xác suất theo công thức p_i = exp(z_i) / sum_j
exp(z_j). Nhãn dự đoán y_hat được chọn bằng argmax(p_i). Trong triển khai
online, gateway không chỉ dùng y_hat mà còn dùng max(p_i) làm confidence. Nếu
confidence nhỏ hơn ngưỡng, frame đó được xem là không đủ chắc chắn để phát lệnh.

Accuracy được tính bằng số mẫu dự đoán đúng chia cho tổng số mẫu test. Precision
của một lớp đo tỷ lệ dự đoán thuộc lớp đó thật sự đúng; recall đo tỷ lệ mẫu thật
của lớp đó được nhận ra đúng; F1 là trung bình điều hòa giữa precision và recall.
Macro F1 lấy trung bình F1 của tất cả các lớp, không phụ thuộc số mẫu từng lớp.
Vì dataset có nhiều lớp điều khiển với ý nghĩa an toàn khác nhau, macro F1 phù
hợp hơn accuracy khi cần đánh giá cân bằng giữa các lớp.

Với latency, báo cáo dùng median và p95 thay vì chỉ dùng trung bình. Median cho
biết trải nghiệm thường gặp, còn p95 phản ánh các trường hợp chậm nhưng vẫn có
thể xuất hiện trong demo. Nếu p95 quá cao, người dùng sẽ cảm thấy hệ thống có
lúc bị khựng dù median vẫn đẹp. Do đó log online cần ghi riêng cloud_rtt_ms,
inference_ms, esp32_ack_ms và total_ms. Cách đo tách đoạn giúp xác định nghẽn ở
cloud, gateway hay đường truyền tới ESP32.

Thuật toán ổn định quyết định có thể mô tả như sau: với mỗi frame, gateway lấy
nhãn và confidence từ cloud; nếu confidence thấp, đưa frame vào trạng thái
ignored; nếu confidence đủ cao, đưa nhãn vào cửa sổ voting có kích thước k; nếu
nhãn chiếm đa số trong cửa sổ và khác lệnh hiện tại, gateway kiểm tra cooldown;
nếu cooldown đã hết và TTL còn hợp lệ, gateway tạo bản tin điều khiển mới. Khi
không có frame hợp lệ trong thời gian deadman_timeout_ms, gateway hoặc firmware
đưa xe về lệnh stop. Thuật toán này làm giảm số lần đổi lệnh đột ngột do một
frame nhiễu đơn lẻ.

| **Đại lượng** | **Công thức/cách tính** | **Vai trò trong đề tài** |
|:---|:---|:---|
| Softmax | p_i = exp(z_i) / sum_j exp(z_j) | Chuyển logit thành confidence |
| Accuracy | correct / total | Đánh giá tổng quát trên test set |
| Precision | TP / (TP + FP) | Đo mức tin cậy khi model dự đoán một lớp |
| Recall | TP / (TP + FN) | Đo khả năng nhận ra mẫu thật của lớp |
| F1 | 2PR / (P + R) | Cân bằng precision và recall |
| Macro F1 | mean(F1_c) với mọi lớp c | So sánh hiệu năng giữa các lớp |
| p95 latency | Phân vị 95% của latency | Đánh giá độ trễ xấu nhưng thường gặp |
| Voting window | majority(label trong k frame) | Ổn định quyết định online |

*Bảng 3.5. Công thức và đại lượng đánh giá sử dụng trong báo cáo*
"""
    ch4_extra = """
## **4.11 Tiêu chí nghiệm thu hệ thống**

Ngoài accuracy và macro F1, hệ thống điều khiển bằng cử chỉ cần được nghiệm thu
theo tiêu chí vận hành. Một mô hình đạt điểm cao trên test set nhưng tạo lệnh
không ổn định trong demo vẫn chưa đạt yêu cầu của đề tài. Vì vậy nhóm dùng thêm
các tiêu chí như tỷ lệ frame bị chặn do confidence thấp, số lần đổi lệnh ngoài
ý muốn, tỷ lệ ACK từ ESP32, số lần deadman timeout kích hoạt và độ trễ p95. Các
tiêu chí này phản ánh trải nghiệm điều khiển thực tế tốt hơn so với chỉ nhìn
vào accuracy.

| **Tiêu chí** | **Cách đo** | **Ngưỡng mong muốn** | **Ý nghĩa** |
|:---|:---|:---|:---|
| Macro F1 offline | Test set theo subject split | >= 0.90 trong bản cải tiến | Đo cân bằng giữa các lớp |
| Cloud RTT p95 | Log gateway khi gọi Azure | < 250 ms | Đảm bảo phản hồi online |
| ACK ESP32 | Tỷ lệ lệnh có phản hồi | > 95% trong mạng ổn định | Đo độ tin cậy truyền lệnh |
| False activation | Số lệnh sai khi không có cử chỉ | Càng thấp càng tốt | Đánh giá an toàn |
| Deadman timeout | Số lần tự dừng khi mất tín hiệu | Phải kích hoạt đúng | Bảo vệ phần cứng |
| Servo cooldown | Khoảng cách giữa hai lệnh servo | Không làm tay robot giật | Giảm rủi ro cơ khí |

*Bảng 4.4. Tiêu chí nghiệm thu hệ thống online*

## **4.12 Kịch bản kiểm thử demo**

Kịch bản kiểm thử online nên được chia thành ba nhóm. Nhóm thứ nhất là kiểm thử
từng cử chỉ trong điều kiện nền đơn giản, ánh sáng ổn định và khoảng cách camera
cố định. Mục tiêu của nhóm này là xác nhận pipeline hoạt động từ camera đến
ESP32, không nhằm chứng minh độ robust. Mỗi cử chỉ được giữ trong 5-10 giây,
gateway ghi log nhãn dự đoán, confidence, lệnh phát ra và ACK từ ESP32.

Nhóm thứ hai là kiểm thử chuyển trạng thái. Người dùng lần lượt đổi từ stop
sang forward, từ forward sang left/right, sau đó quay về stop hoặc no_gesture.
Kịch bản này kiểm tra voting, cooldown và deadman timeout. Nếu gateway đổi lệnh
quá nhanh, xe sẽ phản hồi giật. Nếu gateway quá chậm, trải nghiệm điều khiển
không tự nhiên. Do đó log cần ghi cả thời điểm frame, thời điểm cloud phản hồi
và thời điểm nhận ACK.

Nhóm thứ ba là kiểm thử lỗi có chủ đích. Nhóm có thể tắt mạng trong vài giây,
che camera, đưa tay ra khỏi khung hình, hoặc làm cloud trả lỗi để xác nhận hệ
thống tự dừng. Đây là phần thường bị bỏ qua trong báo cáo nhận dạng ảnh, nhưng
lại cần thiết với IoT vì đầu ra của hệ thống tác động trực tiếp lên motor và
servo. Nếu các lỗi này không được kiểm thử, demo có thể chạy tốt trong điều kiện
lý tưởng nhưng mất an toàn khi môi trường thay đổi.

## **4.13 Kế hoạch đánh giá 70/20/10 và 80/20**

Kết quả chính hiện tại dùng subject split, trong đó test set chứa người không
nằm trong train set. Đây là cách đánh giá thận trọng vì nó kiểm tra khả năng
tổng quát hóa sang người mới. Tuy nhiên, để đáp ứng đúng gợi ý của giảng viên,
nhóm nên chạy thêm hai giao thức bổ sung trên Colab hoặc Kaggle. Giao thức
70/20/10 chia dữ liệu theo clip_id, giữ tỷ lệ lớp tương đối đều giữa train,
validation và test. Giao thức 80/20 có thể dùng train/test trực tiếp, hoặc lấy
validation từ một phần train thông qua validation_split.

Khi chạy 70/20/10, cần tránh chia từng frame ngẫu nhiên nếu các frame trong cùng
clip rất giống nhau. Nếu frame của cùng một clip xuất hiện cả trong train và
test, kết quả sẽ lạc quan hơn thực tế. Cách phù hợp hơn là gom theo clip_id,
sau đó stratified theo gesture. Với 570 clip, tỷ lệ 70/20/10 tương ứng khoảng
399 clip train, 114 clip test và 57 clip validation. Sau khi chia clip, toàn bộ
frame thuộc cùng clip phải đi theo một split duy nhất.

Khi chạy 80/20, nhóm nên báo cáo rõ mục tiêu của giao thức. Nếu 80/20 theo clip
ngẫu nhiên, kết quả trả lời câu hỏi mô hình học tốt trên phân phối dữ liệu hiện
có hay không. Nếu 80/20 theo subject-aware holdout, kết quả gần với subject split
hơn. Hai giao thức không thay thế nhau; chúng bổ sung góc nhìn khác nhau. Báo
cáo cuối nên giữ subject split làm kết quả chính và đưa 70/20/10, 80/20 như thực
nghiệm bổ sung nếu chạy kịp.

## **4.14 Phân tích lỗi và hướng cải thiện dữ liệu**

Nhóm lớp one, two, like và dislike cần được ưu tiên cải thiện vì các lớp này có
hình dạng bàn tay dễ chồng lấn khi crop không đủ rộng hoặc góc cổ tay thay đổi.
Với one và two, sự khác biệt chính nằm ở số ngón giơ lên; nếu đầu ngón bị mờ
hoặc nằm sát mép ảnh, CNN có thể học đặc trưng nền hoặc cổ tay thay vì cấu trúc
ngón. Với like và dislike, hướng ngón cái phụ thuộc mạnh vào góc xoay bàn tay,
nên dữ liệu cần có nhiều hướng quay hơn thay vì chỉ một tư thế chuẩn.

Để cải thiện dữ liệu, nhóm nên thu thêm subject s06-s10, mỗi người thực hiện đủ
tám lớp trong ba điều kiện: nền đơn giản, nền phức tạp và ánh sáng yếu. Mỗi clip
nên ghi metadata background, lighting, handedness và khoảng cách camera. Nếu có
thời gian, nhóm nên thu thêm lớp no_gesture trong tình huống thật, ví dụ người đi
ngang camera, bàn tay cầm vật khác hoặc khung hình chỉ có nền. Các mẫu âm tính
này giúp giảm false activation trong demo.

Về mô hình, hướng cải thiện ngắn hạn là fine-tune thêm các block cuối của
MobileNetV3Small, tăng augmentation có kiểm soát và thử focal loss nếu một số
lớp tiếp tục yếu. Hướng cải thiện dài hạn là kết hợp landmark sequence hoặc
CNN-LSTM cho cử chỉ động. Tuy nhiên, mô hình phức tạp hơn chỉ có ý nghĩa khi có
dữ liệu chuỗi đủ tốt; nếu dataset vẫn nhỏ, tăng kiến trúc có thể làm overfit
nhanh hơn thay vì cải thiện demo.

## **4.15 Bảng phân tích lỗi và hành động khắc phục**

Phân tích lỗi cần đi xa hơn việc đọc một con số accuracy tổng. Với hệ điều khiển
robot, mỗi lỗi có mức độ rủi ro khác nhau. Nhầm giữa hai cử chỉ cùng nhóm servo
có thể chỉ làm tay robot đổi góc sai, nhưng nhầm từ stop sang forward có thể làm
xe di chuyển ngoài ý muốn. Vì vậy báo cáo cần phân loại lỗi theo nguồn gốc và
đề xuất hành động khắc phục tương ứng.

| **Nhóm lỗi** | **Nguyên nhân có khả năng** | **Ảnh hưởng vận hành** | **Hướng khắc phục** |
|:---|:---|:---|:---|
| one nhầm sang two | Đầu ngón bị mờ, crop mất một phần bàn tay | Lệnh có thể bị ánh xạ sai nếu hai lớp điều khiển khác nhau | Thu thêm dữ liệu nhiều góc tay, tăng crop padding, dùng landmark hỗ trợ |
| like nhầm sang dislike | Hướng ngón cái phụ thuộc góc xoay cổ tay | Tay robot hoặc xe có thể phản hồi ngược ý định | Tăng augmentation rotation, thu thêm mẫu tay trái/tay phải |
| stop nhầm sang lớp di chuyển | Background hoặc bàn tay bị che một phần | Rủi ro an toàn cao nhất | Tăng ngưỡng confidence cho lệnh di chuyển, yêu cầu voting dài hơn |
| no_gesture nhầm thành gesture | Vật thể nền giống bàn tay hoặc người đi ngang | False activation khi người dùng chưa ra lệnh | Thu thêm mẫu âm tính, kiểm thử no_gesture trong nhiều bối cảnh |
| Latency tăng đột biến | Cold start cloud, Wi-Fi yếu, request queue | Robot phản hồi chậm hoặc không đều | Warm-up endpoint, log p95, cân nhắc local fallback |
| ACK ESP32 mất | WebSocket mất kết nối hoặc ESP32 reset | Gateway không biết lệnh đã thực thi chưa | Reconnect, sequence number, deadman timeout |

*Bảng 4.5. Phân tích lỗi chính và hướng khắc phục*

Từ bảng phân tích lỗi có thể thấy cải thiện model và cải thiện hệ thống phải đi
song song. Nếu chỉ tăng accuracy nhưng không thay đổi safety filter, một số lỗi
hiếm vẫn có thể gây chuyển động không mong muốn. Ngược lại, nếu chỉ tăng ngưỡng
confidence quá cao, hệ thống sẽ an toàn hơn nhưng kém nhạy và người dùng phải
giữ cử chỉ lâu. Điểm cân bằng cần được chọn dựa trên log thực nghiệm, không nên
chọn thủ công một lần rồi giữ cố định.

Trong bản mở rộng, nhóm nên tạo một dashboard nhỏ đọc trực tiếp từ CSV log để
hiển thị số frame theo nhãn, confidence trung bình, tỷ lệ bị chặn, số lệnh đã
phát, số ACK nhận được và phân bố latency. Dashboard này giúp quá trình demo
minh bạch hơn: khi xe không chạy, nhóm có thể biết nguyên nhân là model không
tự tin, cloud chậm, WebSocket lỗi hay firmware đang trong cooldown. Đây cũng là
cách biến log kỹ thuật thành bằng chứng thuyết phục trong phần bảo vệ.
"""
    body = before_refs[:ch3] + "\n" + source_check + "\n" + before_refs[ch3:]
    body = body.replace("# **CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ THẢO LUẬN**", ch3_extra + "\n# **CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ THẢO LUẬN**")
    body = body.replace("# **KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN**", ch4_extra + "\n# **KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN**")
    body = body.replace("\nPhụ lục\n", "\n")
    body = body.replace("Hình 3.3. Ảnh mẫu gesture\n", "Hình 3.3. Ảnh mẫu gesture\n\nHình 3.4. Ảnh xe tự hành và tay robot sau khi lắp ráp\n\nHình 3.5. Cụm tay robot servo trong quá trình lắp ráp\n")
    body = body.replace("Bảng 2.1. So sánh nghiên cứu liên quan\n", "Bảng 2.1. So sánh nghiên cứu liên quan\n\nBảng 2.3. Ma trận nguồn khoa học và vai trò sử dụng trong báo cáo\n")
    body = body.replace("Bảng 3.2. Tham số model\n", "Bảng 3.2. Tham số model\n\nBảng 3.3. Rủi ro phần cứng và biện pháp giảm thiểu\n\nBảng 3.4. Schema bản tin điều khiển gateway - ESP32\n\nBảng 3.5. Công thức và đại lượng đánh giá sử dụng trong báo cáo\n")
    body = body.replace("Bảng 4.3. Độ trễ online\n", "Bảng 4.3. Độ trễ online\n\nBảng 4.4. Tiêu chí nghiệm thu hệ thống online\n\nBảng 4.5. Phân tích lỗi chính và hướng khắc phục\n")
    body = body.replace(
        "*Hình 3.3. Một số ảnh mẫu trong bộ dữ liệu cử chỉ*\n",
        "*Hình 3.3. Một số ảnh mẫu trong bộ dữ liệu cử chỉ*\n\n"
        "Hình 3.4 minh họa prototype phần cứng sau khi lắp ráp, gồm khung xe, tay robot servo, cụm pin và các mạch điều khiển đặt trên thân xe. Ảnh này cho thấy mô hình đề xuất không chỉ dừng ở mức mô phỏng phần mềm mà đã được tích hợp với cơ cấu chấp hành thật.\n\n"
        '<img src="lab Robot/vehicle_arm_report.png"\n'
        '     alt="Xe tự hành và tay robot sau khi lắp ráp"\n'
        '     width="720" />\n\n'
        "*Hình 3.4. Ảnh xe tự hành và tay robot sau khi lắp ráp*\n",
    )
    return body + "\n" + refs


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.match(r"^\|?\s*:?-{3,}", line):
            i += 1
            continue
        cells = [clean_inline(c.strip()) for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    caption = None
    while i < len(lines) and not lines[i].strip():
        i += 1
    if i < len(lines) and re.match(r"^\*Bảng\s+", lines[i].strip()):
        caption = clean_inline(lines[i].strip())
        i += 1
    return rows, caption, i


def add_md(doc, md):
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith(">"):
            i += 1
            continue
        if line.startswith("# "):
            heading_text = clean_inline(line[2:]).upper()
            if heading_text.startswith(("CHƯƠNG ", "KẾT LUẬN", "TÀI LIỆU THAM KHẢO")):
                doc.add_page_break()
            add_heading(doc, line[2:], 1)
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:], 2)
            i += 1
            continue
        if line.startswith("### "):
            add_heading(doc, line[4:], 3)
            i += 1
            continue
        if line.startswith("|"):
            rows, caption, i = parse_table(lines, i)
            if caption:
                add_table_caption(doc, caption)
            add_table(doc, rows)
            continue
        if line.startswith("<img src="):
            m = re.search(r'<img src="([^"]+)"', line)
            src = m.group(1) if m else ""
            while i < len(lines) and "/>" not in lines[i]:
                i += 1
            i += 1
            while i < len(lines) and not lines[i].strip():
                i += 1
            caption = ""
            if i < len(lines) and re.match(r"^\*Hình\s+", lines[i].strip()):
                caption = clean_inline(lines[i].strip())
                i += 1
            add_figure(doc, src, caption or "Hình minh họa")
            continue
        if line.startswith("*Hình "):
            i += 1
            continue
        if line.startswith("*Bảng "):
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_para_format(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False, after=3)
            r = p.add_run(clean_inline(line[2:]))
            set_run_font(r, size=11)
            i += 1
            continue

        para = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or nxt.startswith(("#", "|", "<img", "*Hình ", "*Bảng ", "- ", ">")):
                break
            para.append(nxt)
            i += 1
        add_p(doc, " ".join(para))


def add_break_before_main(doc):
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)
    section.footer.is_linked_to_previous = True


def main():
    md = SOURCE_MD.read_text(encoding="utf-8")
    doc = Document()
    setup_document(doc)
    cover(doc)
    add_md(doc, transform_markdown(md))
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()

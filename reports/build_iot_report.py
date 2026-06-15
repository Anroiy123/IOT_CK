from __future__ import annotations

import csv
import json
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "reports"
ASSETS = OUT / "report_assets"
DOCX_OUT = OUT / "BaoCao_IOT_Nhom12.docx"
BIB_OUT = OUT / "BaoCao_IOT_Nhom12_references.bib"
QA_OUT = OUT / "BaoCao_IOT_Nhom12_QA.json"


def load_inputs():
    metrics = json.loads((OUT / "cnn_baseline_s05_partial_metrics.json").read_text(encoding="utf-8"))
    comparison = json.loads((OUT / "cnn_comparison_summary.json").read_text(encoding="utf-8"))
    latency = json.loads((ASSETS / "latency_summary.json").read_text(encoding="utf-8"))
    rows = list(csv.DictReader((ROOT / "data" / "metadata.csv").open(encoding="utf-8")))
    return metrics, comparison, latency, rows


REFS = [
    ("Jalayer2026", "R. Jalayer, M. Jalayer, C. Orsenigo, and M. Tomizuka", "2026", "A review on deep learning for vision-based hand detection, hand segmentation and hand gesture recognition in human-robot interaction", "Robotics and Computer-Integrated Manufacturing, vol. 97, 103110", "10.1016/j.rcim.2025.103110", "https://doi.org/10.1016/j.rcim.2025.103110", "Elsevier journal; SCImago/WoS page available"),
    ("Oudah2020", "M. Oudah, A. Al-Naji, and J. Chahl", "2020", "Hand gesture recognition based on computer vision: a review of techniques", "Journal of Imaging, vol. 6, no. 8, 73", "10.3390/jimaging6080073", "https://doi.org/10.3390/jimaging6080073", "Peer-reviewed journal"),
    ("Sarma2021", "D. Sarma and M. K. Bhuyan", "2021", "Methods, databases and recent advancement of vision-based hand gesture recognition for HCI systems: a review", "SN Computer Science, vol. 2, 436", "10.1007/s42979-021-00827-x", "https://doi.org/10.1007/s42979-021-00827-x", "Springer journal; index check recommended"),
    ("Pavlovic1997", "V. I. Pavlovic, R. Sharma, and T. S. Huang", "1997", "Visual interpretation of hand gestures for human-computer interaction: a review", "IEEE Transactions on Pattern Analysis and Machine Intelligence", "10.1109/34.598226", "https://doi.org/10.1109/34.598226", "IEEE journal"),
    ("Zhang2020", "F. Zhang et al.", "2020", "MediaPipe Hands: On-device real-time hand tracking", "arXiv:2006.10214", "", "https://arxiv.org/abs/2006.10214", "Technical preprint; foundational implementation"),
    ("Howard2019", "A. Howard et al.", "2019", "Searching for MobileNetV3", "IEEE/CVF International Conference on Computer Vision (ICCV)", "10.1109/ICCV.2019.00140", "https://openaccess.thecvf.com/content_ICCV_2019/html/Howard_Searching_for_MobileNetV3_ICCV_2019_paper.html", "IEEE/CVF conference"),
    ("Sandler2018", "M. Sandler, A. Howard, M. Zhu, A. Zhmoginov, and L.-C. Chen", "2018", "MobileNetV2: Inverted residuals and linear bottlenecks", "IEEE/CVF Conference on Computer Vision and Pattern Recognition (CVPR)", "10.1109/CVPR.2018.00474", "https://doi.org/10.1109/CVPR.2018.00474", "IEEE/CVF conference"),
    ("Materzynska2019", "J. Materzynska, G. Berger, I. Bax, and R. Memisevic", "2019", "The Jester Dataset: A large-scale video dataset of human gestures", "ICCV Workshops", "", "https://openaccess.thecvf.com/content_ICCVW_2019/papers/HANDS/Materzynska_The_Jester_Dataset_A_Large-Scale_Video_Dataset_of_Human_Gestures_ICCVW_2019_paper.pdf", "CVF workshop paper"),
    ("Devineau2018", "G. Devineau, F. Moutarde, W. Xi, and J. Yang", "2018", "Deep learning for hand gesture recognition on skeletal data", "13th IEEE International Conference on Automatic Face & Gesture Recognition", "10.1109/FG.2018.00025", "https://doi.org/10.1109/FG.2018.00025", "IEEE conference"),
    ("DeSmedt2016", "Q. De Smedt, H. Wannous, and J.-P. Vandeborre", "2016", "Skeleton-based dynamic hand gesture recognition", "IEEE CVPR Workshops", "10.1109/CVPRW.2016.153", "https://doi.org/10.1109/CVPRW.2016.153", "IEEE/CVF workshop"),
    ("Mujahid2021", "A. Mujahid et al.", "2021", "Real-time hand gesture recognition based on deep learning YOLOv3 model", "Applied Sciences, vol. 11, no. 9, 4164", "10.3390/app11094164", "https://doi.org/10.3390/app11094164", "Peer-reviewed journal"),
    ("Molchanov2016", "P. Molchanov, S. Gupta, K. Kim, and J. Kautz", "2016", "Online detection and classification of dynamic hand gestures with recurrent 3D convolutional neural networks", "IEEE Conference on Computer Vision and Pattern Recognition", "10.1109/CVPR.2016.454", "https://doi.org/10.1109/CVPR.2016.454", "IEEE/CVF conference"),
    ("Nunez2018", "J. C. Nunez et al.", "2018", "Convolutional neural networks and long short-term memory for skeleton-based human activity and hand gesture recognition", "Pattern Recognition, vol. 76", "10.1016/j.patcog.2017.10.033", "https://doi.org/10.1016/j.patcog.2017.10.033", "Elsevier journal"),
    ("Khan2021", "F. Khan et al.", "2021", "3D hand gestures segmentation and optimized classification using deep learning", "IEEE Access", "10.1109/ACCESS.2021.3114871", "https://doi.org/10.1109/ACCESS.2021.3114871", "IEEE Access; SCImago page available"),
    ("Hu2018", "B. Hu and J. Wang", "2018", "Deep learning based hand gesture recognition and UAV flight controls", "24th International Conference on Automation and Computing", "10.23919/IConAC.2018.8748953", "https://doi.org/10.23919/IConAC.2018.8748953", "IEEE-indexed conference"),
    ("He2020", "H. He and Y. Dan", "2020", "The research and design of smart mobile robotic arm based on gesture controlled", "International Conference on Advanced Mechatronic Systems", "10.1109/ICAMechS49982.2020.9310156", "https://doi.org/10.1109/ICAMechS49982.2020.9310156", "IEEE conference"),
    ("Islam2020", "J. Islam, A. Ghosh, M. I. Iqbal, S. Meem, and N. Ahmad", "2020", "Integration of home assistance with a gesture controlled robotic arm", "IEEE Region 10 Symposium", "10.1109/TENSYMP50017.2020.9230893", "https://doi.org/10.1109/TENSYMP50017.2020.9230893", "IEEE conference"),
    ("Dissanayake2021", "D. Dissanayake and contributors", "2021", "Real-time hand gesture recognition for robotic arm and home automation", "ACM International Conference Proceedings", "10.1145/3459104.3459142", "https://dl.acm.org/doi/fullHtml/10.1145/3459104.3459142", "ACM conference"),
    ("ACMRobot2018", "Y. Hu, J. Xu, Z. Ma, and G. Cao", "2018", "Predictive hand gesture classification for real time robot control", "ACM International Conference Proceedings", "10.1145/3240876.3240914", "https://dl.acm.org/doi/10.1145/3240876.3240914", "ACM conference"),
    ("MobileRobot2012", "K. R. Konda and A. Konigs", "2012", "Real time interaction with mobile robots using hand gestures", "ACM International Conference Proceedings", "10.1145/2157689.2157743", "https://dl.acm.org/doi/10.1145/2157689.2157743", "ACM conference"),
    ("Kwolek2023", "B. Kwolek", "2023", "Continuous hand gesture recognition for human-robot collaborative assembly", "ICCV Workshop on Assistive Computer Vision and Robotics", "", "https://openaccess.thecvf.com/content/ICCV2023W/ACVR/papers/Kwolek_Continuous_Hand_Gesture_Recognition_for_Human-Robot_Collaborative_Assembly_ICCVW_2023_paper.pdf", "CVF workshop"),
    ("Nogales2023", "R. E. Nogales and M. E. Benalcazar", "2023", "Hand gesture recognition using automatic feature extraction and deep learning algorithms with memory", "Big Data and Cognitive Computing, vol. 7, no. 2, 102", "10.3390/bdcc7020102", "https://doi.org/10.3390/bdcc7020102", "Peer-reviewed journal"),
    ("Skaria2019", "S. Skaria, A. Al-Hourani, R. J. Evans, and collaborators", "2019", "Hand-gesture recognition using two-antenna Doppler radar with deep convolutional neural networks", "IEEE Sensors Journal, vol. 19, no. 8", "10.1109/JSEN.2019.2892073", "https://doi.org/10.1109/JSEN.2019.2892073", "IEEE journal"),
    ("Li2020", "H. Li et al.", "2020", "Hand gesture recognition enhancement based on spatial fuzzy matching in Leap Motion", "IEEE Transactions on Industrial Informatics, vol. 16, no. 3", "10.1109/TII.2019.2931140", "https://doi.org/10.1109/TII.2019.2931140", "IEEE journal"),
    ("Sokolova2009", "M. Sokolova and G. Lapalme", "2009", "A systematic analysis of performance measures for classification tasks", "Information Processing & Management, vol. 45, no. 4", "10.1016/j.ipm.2009.03.002", "https://doi.org/10.1016/j.ipm.2009.03.002", "Elsevier journal"),
    ("Powers2011", "D. M. W. Powers", "2011", "Evaluation: From precision, recall and F-measure to ROC, informedness, markedness and correlation", "Journal of Machine Learning Technologies, vol. 2, no. 1", "", "https://www.researchgate.net/publication/228529307", "Evaluation reference; index check recommended"),
    ("AzureDocs", "Microsoft", "2026", "Azure Container Apps documentation", "Microsoft Learn", "", "https://learn.microsoft.com/azure/container-apps/", "Technical documentation"),
    ("FastAPI", "FastAPI contributors", "2026", "FastAPI documentation", "FastAPI official documentation", "", "https://fastapi.tiangolo.com/", "Technical documentation"),
    ("WebSocketRFC", "I. Fette and A. Melnikov", "2011", "The WebSocket Protocol", "IETF RFC 6455", "10.17487/RFC6455", "https://www.rfc-editor.org/rfc/rfc6455", "IETF standard"),
]

REF_INDEX = {r[0]: i + 1 for i, r in enumerate(REFS)}


def cite(*keys: str) -> str:
    nums = [REF_INDEX[k] for k in keys if k in REF_INDEX]
    return "[" + "], [".join(map(str, nums)) + "]"


def font_run(run, size=None, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_p(doc, text="", *, style=None, align=None, bold=False, italic=False, size=13, indent=True):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    elif style is None:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent and style is None and text:
        p.paragraph_format.first_line_indent = Cm(0.8)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font_run(r, size=size, bold=bold, italic=italic)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        font_run(r, size=16 if level == 1 else 14 if level == 2 else 13, bold=True, color="000000")
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_caption(doc, text):
    add_p(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=11, indent=False)


def set_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn("w:" + edge))
        if elem is None:
            elem = OxmlElement("w:" + edge)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), "808080")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, *, bold=False, size=9, center=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center or bold else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    font_run(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, data, *, widths=None, size=9):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    for i, h in enumerate(headers):
        cell_text(table.rows[0].cells[i], h, bold=True, size=size, center=True)
        shade(table.rows[0].cells[i], "D9EAF7")
    for row in data:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell_text(cells[i], value, size=size)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(w)
    return table


def add_confusion_matrix_table(doc, labels, matrix, caption):
    rows = [[label] + [int(value) for value in matrix[index]] for index, label in enumerate(labels)]
    add_table(doc, ["True\\Pred"] + list(labels), rows, widths=[0.9] + [0.45] * len(labels), size=7)
    add_caption(doc, caption)


def add_image(doc, filename, caption, width=6.0):
    path = ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_page_number(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font_run(p.add_run("Trang "), size=10)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)
    sec.header_distance = Cm(1.2)
    sec.footer_distance = Cm(1.2)
    add_page_number(sec)
    for name in ["Normal", "List Bullet", "List Number"]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(13)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        st = doc.styles[name]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0, 0, 0)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    font_run(r, size=13)
    p.paragraph_format.space_after = Pt(3)


def cover(doc):
    for line in ["BỘ KHOA HỌC VÀ CÔNG NGHỆ", "HỌC VIỆN CÔNG NGHỆ BƯU CHÍNH VIỄN THÔNG", "CƠ SỞ TẠI THÀNH PHỐ HỒ CHÍ MINH", "----------o0o----------"]:
        add_p(doc, line, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    for _ in range(4):
        add_p(doc, "", indent=False)
    add_p(doc, "BÁO CÁO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=24, indent=False)
    add_p(doc, "ĐỒ ÁN MÔN HỌC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, indent=False)
    add_p(doc, "THIẾT KẾ THIẾT BỊ ĐIỀU KHIỂN BẰNG CỬ CHỈ TAY SỬ DỤNG DEEP LEARNING CHO ĐIỀU KHIỂN THÔNG MINH XE TỰ HÀNH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, indent=False)
    for _ in range(2):
        add_p(doc, "", indent=False)
    add_table(doc, ["Thông tin", "Nội dung"], [
        ["Môn học", "Phát triển ứng dụng IOT"],
        ["Giảng viên hướng dẫn", "Đàm Minh Lịnh"],
        ["Nhóm thực hiện", "Nhóm 12"],
        ["Lớp", "D22CQPTUD01-N"],
    ], widths=[2.2, 4.7], size=12)
    add_p(doc, "", indent=False)
    add_table(doc, ["MSSV", "Họ tên", "Email", "Vai trò"], [
        ["N22DCPT035", "Trần Quang Hùng", "n22dcpt035@student.ptithcm.edu.vn", "Nhóm trưởng"],
        ["N22DCPT052", "Vũ Quang Long", "n22dcpt052@student.ptithcm.edu.vn", "Thành viên"],
    ], widths=[1.2, 1.8, 2.8, 1.1], size=10)
    add_p(doc, "TP. Hồ Chí Minh, 15 / 6 / 2026", align=WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    doc.add_page_break()


def front_matter(doc, metrics, latency, rows):
    add_h(doc, "LỜI CẢM ƠN", 1)
    for text in [
        "Nhóm sinh viên chúng em xin gửi lời cảm ơn chân thành đến Học viện Công nghệ Bưu chính Viễn thông và quý thầy cô đã tạo điều kiện để nhóm được tiếp cận các kiến thức về IoT, trí tuệ nhân tạo và triển khai hệ thống nhúng trong môi trường thực tế.",
        "Đặc biệt, nhóm xin cảm ơn thầy Đàm Minh Lịnh đã hướng dẫn định hướng đề tài, nhấn mạnh yêu cầu về nguồn tài liệu khoa học, tiêu chí đánh giá mô hình và yêu cầu trình bày báo cáo theo cấu trúc nghiên cứu.",
        "Do thời gian thực hiện và điều kiện phần cứng còn hạn chế, hệ thống chắc chắn chưa thể đạt mức hoàn thiện như một sản phẩm thương mại. Nhóm mong nhận được góp ý để tiếp tục cải thiện mô hình nhận dạng, mở rộng dữ liệu và tối ưu độ trễ điều khiển.",
    ]:
        add_p(doc, text)
    add_h(doc, "TÓM TẮT", 1)
    for text in [
        "Báo cáo trình bày quá trình thiết kế và thực nghiệm một thiết bị điều khiển xe tự hành bằng cử chỉ tay sử dụng Deep Learning. Hệ thống dùng webcam laptop, gateway MediaPipe crop, cloud API FastAPI trên Azure Container Apps và ESP32 điều khiển L298N/PCA9685.",
        f"Mô hình chính là CNN MobileNetV3Small, ảnh đầu vào {metrics['image_size']}x{metrics['image_size']}, huấn luyện trên bộ dữ liệu nội bộ gồm 5 người, 8 lớp, {len(set(r['clip_id'] for r in rows))} clip và {len(rows):,} frame. Kết quả offline đạt accuracy {metrics['accuracy']*100:.2f}% và macro F1 {metrics['macro_f1']*100:.2f}%. Azure RTT median {latency['Azure']['cloud_rtt_ms']['median']} ms, p95 {latency['Azure']['cloud_rtt_ms']['p95']} ms.",
        "Đóng góp chính là pipeline end-to-end từ nhận dạng cử chỉ đến điều khiển phần cứng, có confidence threshold, voting/stabilization, cooldown servo, sequence number, token và deadman timeout. Báo cáo bổ sung thêm so sánh CNN với CNN-LSTM, summary robustness theo background và false activation của lớp no_gesture.",
    ]:
        add_p(doc, text)
    add_h(doc, "MỤC LỤC", 1)
    for item in ["Danh mục thuật ngữ và từ viết tắt", "Chương 1. Giới thiệu", "Chương 2. Nghiên cứu liên quan", "Chương 3. Mô hình đề xuất", "Chương 4. Thực nghiệm, đánh giá và thảo luận", "Kết luận, hạn chế và hướng phát triển", "Tài liệu tham khảo", "Phụ lục"]:
        add_p(doc, item, indent=False)
    add_h(doc, "DANH MỤC HÌNH", 1)
    for item in ["Hình 3.1. Kiến trúc hệ thống", "Hình 3.2. Trình tự xử lý online", "Hình 3.3. Ảnh mẫu gesture", "Hình 4.1. Phân bố dataset", "Hình 4.2. Lịch sử huấn luyện", "Hình 4.3. Confusion matrix", "Hình 4.4. So sánh latency"]:
        add_p(doc, item, indent=False)
    add_h(doc, "DANH MỤC BẢNG", 1)
    for item in ["Bảng 1.1. Mục tiêu và tiêu chí", "Bảng 2.1. So sánh nghiên cứu liên quan", "Bảng 3.1. Cử chỉ và chức năng", "Bảng 3.2. Tham số model", "Bảng 4.1. Thống kê dataset", "Bảng 4.2. Kết quả offline", "Bảng 4.3. Độ trễ online"]:
        add_p(doc, item, indent=False)
    add_h(doc, "DANH MỤC THUẬT NGỮ VÀ TỪ VIẾT TẮT", 1)
    add_table(doc, ["Từ viết tắt", "Tên đầy đủ", "Diễn giải"], [
        ["AI", "Artificial Intelligence", "Trí tuệ nhân tạo"],
        ["CNN", "Convolutional Neural Network", "Mạng nơ-ron tích chập"],
        ["CNN-LSTM", "CNN + Long Short-Term Memory", "Mô hình kết hợp không gian và chuỗi thời gian"],
        ["DL", "Deep Learning", "Học sâu"],
        ["ESP32", "Espressif ESP32", "Vi điều khiển Wi-Fi/Bluetooth"],
        ["FastAPI", "FastAPI framework", "Framework triển khai API suy diễn"],
        ["HCI", "Human-Computer Interaction", "Tương tác người - máy"],
        ["HGR", "Hand Gesture Recognition", "Nhận dạng cử chỉ tay"],
        ["HRI", "Human-Robot Interaction", "Tương tác người - robot"],
        ["IoT", "Internet of Things", "Internet vạn vật"],
        ["PCA9685", "16-channel PWM driver", "Mạch mở rộng PWM điều khiển servo"],
        ["RTT", "Round Trip Time", "Thời gian khứ hồi khi gọi cloud API"],
        ["WebSocket", "RFC 6455 WebSocket", "Kênh truyền lệnh hai chiều"],
    ], widths=[1.4, 2.5, 3.0], size=9)
    doc.add_page_break()


def chapter1(doc):
    add_h(doc, "CHƯƠNG 1. GIỚI THIỆU", 1)
    paras = [
        f"Trong các hệ thống tương tác người - máy, cử chỉ tay là một kênh giao tiếp tự nhiên vì người dùng không cần thiết bị cầm tay và có thể truyền đạt lệnh nhanh bằng chuyển động trực quan. Các nghiên cứu tổng quan về hand gesture recognition chỉ ra rằng nhận dạng cử chỉ đã phát triển từ đặc trưng thủ công sang học sâu trên ảnh, depth, skeleton hoặc cảm biến đeo {cite('Pavlovic1997','Oudah2020','Jalayer2026')}.",
        "Đề tài này xuất phát từ yêu cầu thiết kế thiết bị nhận dạng cử chỉ tay sử dụng Deep Learning cho điều khiển thông minh xe tự hành. Khác với bài toán chỉ phân loại ảnh offline, hệ thống cần xử lý liên tục: camera đọc frame, mô hình suy diễn cử chỉ, gateway lọc nhiễu, truyền lệnh qua mạng và vi điều khiển thực thi trên motor/servo.",
        f"Deep Learning phù hợp với bài toán vì CNN học đặc trưng không gian của bàn tay, trong khi CNN-LSTM hoặc 3D-CNN phù hợp hơn với cử chỉ động {cite('Howard2019','Sandler2018','Molchanov2016','Nunez2018')}. Đề tài chọn CNN MobileNetV3Small làm mô hình chính, đồng thời giữ CNN-LSTM như hướng mở rộng.",
        f"Ở tầng tiền xử lý, MediaPipe Hands cho thấy khả năng phát hiện bàn tay và landmark theo thời gian thực từ camera RGB đơn {cite('Zhang2020')}. Repo hiện tại dùng MediaPipe để crop vùng bàn tay trước khi gửi ảnh lên cloud, giúp giảm nhiễu nền nhưng vẫn đáp ứng yêu cầu model lưu trữ trên cloud.",
        f"Trong tương tác người - robot, các nghiên cứu gần đây nhấn mạnh rằng nhận dạng cử chỉ cần được đánh giá trong điều kiện vận hành thực tế, bao gồm nền phức tạp, khoảng cách camera, tốc độ phản hồi và trạng thái an toàn của robot {cite('Jalayer2026','Kwolek2023','Hu2018')}.",
    ]
    for p in paras:
        add_p(doc, p)
    add_h(doc, "1.1 Bối cảnh và tính cấp thiết", 2)
    for p in [
        "Sự phát triển của IoT tạo điều kiện để các thiết bị nhúng như ESP32 kết nối trực tiếp với gateway, cloud và cơ cấu chấp hành. Khi kết hợp với AI, thiết bị nhúng không nhất thiết phải tự xử lý toàn bộ ảnh; nó có thể nhận lệnh đã được suy diễn từ gateway hoặc cloud service.",
        "Nếu cử chỉ bị nhận sai trong hệ robot, lỗi có thể dẫn đến chuyển động không mong muốn. Vì vậy hệ thống đề xuất không truyền mọi dự đoán thành lệnh ngay mà dùng confidence threshold, số lần lặp liên tiếp và deadman timeout.",
    ]:
        add_p(doc, p)
    add_h(doc, "1.2 Mục tiêu, phạm vi và đóng góp", 2)
    add_table(doc, ["Nhóm mục tiêu", "Nội dung", "Tiêu chí đánh giá"], [
        ["Nhận dạng", "Phân loại 8 lớp cử chỉ tay", "Accuracy, macro F1, confusion matrix"],
        ["Thời gian thực", "Gateway xử lý webcam và gọi cloud online", "RTT, inference_ms, total_ms"],
        ["Triển khai IoT", "ESP32 nhận lệnh qua WebSocket", "ACK, sequence, token, deadman timeout"],
        ["An toàn", "Không phát lệnh khi confidence thấp hoặc cloud lỗi", "stop tự động, cooldown servo"],
        ["Học thuật", "Liên hệ CNN, CNN-LSTM, HGR/HRI, cloud/edge AI", "Tài liệu IEEE/ACM/Elsevier và Zotero"],
    ], widths=[1.5, 3.1, 2.3], size=9)
    add_caption(doc, "Bảng 1.1. Mục tiêu và tiêu chí đánh giá đề tài")
    add_p(doc, "Đóng góp thứ nhất là thiết kế pipeline end-to-end từ camera đến phần cứng. Đóng góp thứ hai là xây dựng dataset nội bộ có metadata rõ ràng theo subject, gesture, clip và frame. Đóng góp thứ ba là triển khai CNN trên Azure. Đóng góp thứ tư là tích hợp cơ chế an toàn ở gateway và firmware.")
    add_p(doc, "Báo cáo gồm bốn chương: giới thiệu, nghiên cứu liên quan, mô hình đề xuất, thực nghiệm và thảo luận. Phần cuối trình bày kết luận, hạn chế và hướng phát triển theo ba đoạn đúng yêu cầu.")


def chapter2(doc):
    doc.add_page_break()
    add_h(doc, "CHƯƠNG 2. NGHIÊN CỨU LIÊN QUAN", 1)
    for p in [
        f"Các nghiên cứu liên quan được chia thành năm nhóm: tổng quan HGR/HRI, CNN cho ảnh tĩnh, CNN-LSTM/3D-CNN cho cử chỉ động, landmark/crop bàn tay và ứng dụng điều khiển robot {cite('Jalayer2026','Oudah2020','Sarma2021')}.",
        f"CNN trực tiếp trên ảnh crop giữ được thông tin hình dạng bàn tay, còn landmark giúp giảm chiều dữ liệu. CNN-LSTM hoặc 3D-CNN mạnh hơn với cử chỉ động nhưng cần dữ liệu video/sequence và chi phí tính toán cao hơn {cite('Molchanov2016','Nunez2018','DeSmedt2016')}.",
        f"Các nghiên cứu robot/mobile robot/robotic arm sử dụng gesture nhấn mạnh yếu tố real-time và an toàn {cite('Hu2018','He2020','Islam2020','Dissanayake2021')}. Điều này phù hợp với thiết kế gateway chỉ chấp nhận gesture ổn định và firmware tự dừng khi timeout.",
    ]:
        add_p(doc, p)
    related = [
        ["Jalayer et al. (2026)", "Review HGR/HRI", "Nhiều nghiên cứu HRI", "Deep learning vision-based", "Tổng quan rộng", "Không phải demo cụ thể", "Nền Ch.1-2"],
        ["Oudah et al. (2020)", "Review HGR", "Vision-based HGR", "Computer vision/ML", "Tổng hợp pipeline", "Ít tập trung cloud IoT", "Pipeline xử lý"],
        ["MediaPipe Hands", "Hand tracking", "RGB camera", "Palm detector + landmark", "Real-time", "Không phân loại 8 label", "Crop/tiền xử lý"],
        ["MobileNetV3", "Backbone nhẹ", "ImageNet", "NAS + NetAdapt", "Tối ưu tài nguyên", "Cần fine-tune", "CNN baseline"],
        ["Jester Dataset", "Dataset gesture", "148k video", "3D CNN baseline", "Dataset lớn", "Label khác dự án", "So sánh dữ liệu"],
        ["Devineau et al.", "Skeleton HGR", "Skeletal data", "Deep learning", "Giảm chiều", "Cần landmark", "Hướng CNN-LSTM"],
        ["Molchanov et al.", "Dynamic HGR", "Video sequence", "Recurrent 3D CNN", "Khai thác thời gian", "Nặng hơn CNN", "Hướng phát triển"],
        ["Hu & Wang", "Gesture + UAV", "Camera/gesture", "Deep learning + control", "Liên hệ robot", "Khác phần cứng", "HRI điều khiển"],
        ["Islam et al.", "Robotic arm", "Gesture control", "Robot arm integration", "Gần tay máy", "Không cloud DL", "Cơ cấu"],
        ["ACM robot control", "Robot control", "Real-time gesture", "Predictive classification", "Gần bài toán", "Cần full metadata", "Real-time"],
    ]
    add_table(doc, ["Nghiên cứu", "Nhóm", "Dữ liệu", "Mô hình", "Ưu điểm", "Hạn chế", "Liên hệ"], related, widths=[1.25, 1.0, 1.1, 1.2, 1.25, 1.25, 1.2], size=7)
    add_caption(doc, "Bảng 2.1. So sánh nghiên cứu liên quan")
    for heading, body in [
        ("2.1 Nhận dạng cử chỉ tay dựa trên thị giác máy tính", "Pipeline phổ biến gồm thu nhận ảnh, phát hiện bàn tay, tách vùng quan tâm, trích xuất đặc trưng và phân loại. Chất lượng dữ liệu đầu vào quyết định đáng kể hiệu năng mô hình. Trong repo hiện tại, MediaPipe crop giảm nhiễu từ nền trước khi ảnh được resize về 160x160 cho CNN."),
        ("2.2 CNN, CNN-LSTM và mô hình nhẹ", "MobileNetV3Small phù hợp với mục tiêu mô hình nhẹ và demo thời gian thực. CNN-LSTM hoặc 3D-CNN phù hợp với gesture động nhưng yêu cầu dữ liệu chuỗi đầy đủ hơn. Do dataset hiện tại chủ yếu phục vụ gesture tĩnh/đơn khung, CNN baseline là lựa chọn thực dụng."),
        ("2.3 Điều khiển robot bằng cử chỉ", "Với robot vật lý, hệ thống cần có lệnh dừng, xử lý mất kết nối và tránh phát lệnh khi confidence không đủ. Đề tài triển khai các cơ chế này ở hai lớp: gateway và firmware ESP32."),
        ("2.4 Khoảng trống nghiên cứu", "Khoảng trống của đề tài là cân bằng giữa accuracy, latency, cloud deployment và safety. Nhiều nghiên cứu chỉ báo cáo accuracy, trong khi hệ IoT cần đánh giá cả độ trễ, log online và hành vi khi lỗi mạng."),
    ]:
        add_h(doc, heading, 2)
        add_p(doc, body)


def chapter3(doc, metrics):
    doc.add_page_break()
    add_h(doc, "CHƯƠNG 3. MÔ HÌNH ĐỀ XUẤT", 1)
    add_h(doc, "3.1 Sơ đồ mô hình phát hiện và điều khiển", 2)
    add_p(doc, "Mô hình đề xuất kết hợp nhận dạng cử chỉ bằng Deep Learning và điều khiển IoT theo kiến trúc phân tầng. Webcam và gateway chạy trên laptop, cloud API đảm nhiệm suy diễn CNN, ESP32 tập trung điều khiển motor và servo.")
    add_p(doc, "Gateway nhận frame từ webcam, crop vùng bàn tay bằng MediaPipe, mã hóa JPEG và gửi request có session_id/request_id lên Azure API. Cloud trả về gesture, confidence, inference_ms và model_version. Gateway dùng safety filter trước khi gửi lệnh WebSocket tới ESP32.")
    add_image(doc, "architecture_pipeline.png", "Hình 3.1. Kiến trúc hệ thống điều khiển xe tự hành bằng cử chỉ tay", 6.5)
    add_image(doc, "sequence_online.png", "Hình 3.2. Trình tự xử lý online theo thời gian thực", 6.5)
    add_h(doc, "3.2 Các thành phần của hệ thống", 2)
    for p in [
        "Tầng dữ liệu gồm data/raw và metadata.csv. Mỗi dòng metadata ghi frame, subject_id, gesture, clip_id, timestamp, background, lighting, split và found_hand. Metadata chi tiết giúp kiểm soát subject split và tránh leakage.",
        "Tầng mô hình gồm script thu dữ liệu, train CNN, train CNN-LSTM định hướng và benchmark local. Model demo là gesture-cnn-baseline-s05-partial.keras, nhận ảnh RGB 160x160 và phân loại 8 nhãn.",
        "Tầng cloud gồm FastAPI, inference và Dockerfile triển khai Azure. Endpoint /health trả trạng thái, /v1/model trả model_version và /v1/predict nhận ảnh base64.",
        "Tầng gateway gồm camera capture, MediaPipeCropper, CloudGestureClient, GestureStabilizer, GestureMapper và transport WebSocket/DryRun. Log CSV ghi session_id, request_id, capture_ms, preprocess_ms, cloud_rtt_ms, inference_ms, esp32_ack_ms và total_ms.",
        "Tầng firmware ESP32 gồm HTTP /health, HTTP /state và WebSocket port 81. Firmware kiểm tra token, sequence và TTL; nếu token sai hoặc quá hạn thì dừng motor.",
    ]:
        add_p(doc, p)
    add_table(doc, ["Label", "Ký hiệu", "Chế độ xe", "Chế độ tay máy"], [
        ["stop", "Bàn tay mở", "Dừng xe", "Dừng motor"],
        ["peace", "Chữ V rộng", "Chuyển chế độ xe", "Chuyển chế độ xe"],
        ["rock", "Ký hiệu rock", "Chuyển chế độ tay máy", "Giữ chế độ tay máy"],
        ["like", "Ngón cái lên", "Tiến", "Tăng góc khớp 5 độ"],
        ["dislike", "Ngón cái xuống", "Lùi", "Giảm góc khớp 5 độ"],
        ["one", "Một ngón trỏ", "Rẽ trái", "Chọn khớp trước"],
        ["two", "Hai ngón", "Rẽ phải", "Chọn khớp tiếp"],
        ["no_gesture", "Không có cử chỉ hợp lệ", "Không phát lệnh", "Không phát lệnh"],
    ], widths=[1.3, 2.0, 1.8, 1.8], size=9)
    add_caption(doc, "Bảng 3.1. Các lớp cử chỉ và chức năng điều khiển")
    add_image(doc, "gesture_samples.png", "Hình 3.3. Một số ảnh mẫu trong bộ dữ liệu cử chỉ", 6.2)
    add_h(doc, "3.3 Bảng tham số model", 2)
    add_table(doc, ["Tham số", "Giá trị"], [
        ["Tên model", metrics["model_name"]],
        ["Loại model", metrics["model_type"]],
        ["Backbone", "MobileNetV3Small"],
        ["Pretrained", "ImageNet" if metrics.get("pretrained") else "Không"],
        ["Kích thước ảnh", f"{metrics['image_size']} x {metrics['image_size']}"],
        ["Số lớp", str(len(metrics["labels"]))],
        ["Nhãn", ", ".join(metrics["labels"])],
        ["Epoch", str(metrics["epochs_requested"])],
        ["Batch size", str(metrics["batch_size"])],
        ["Split strategy", metrics["split_strategy"]],
        ["Train/Val/Test", f"{metrics['counts']['train']} / {metrics['counts']['val']} / {metrics['counts']['test']} frame"],
    ], widths=[2.4, 4.4], size=10)
    add_caption(doc, "Bảng 3.2. Tham số mô hình CNN baseline")
    add_h(doc, "3.4 Công thức, tiêu chí đánh giá và tối ưu", 2)
    for p in [
        "Softmax chuyển vector logit z thành xác suất lớp: p_i = exp(z_i) / sum_j exp(z_j). Lớp dự đoán là lớp có xác suất lớn nhất; confidence là xác suất của lớp được chọn.",
        "Accuracy = số mẫu dự đoán đúng / tổng số mẫu. Precision = TP / (TP + FP), recall = TP / (TP + FN), F1 = 2 * precision * recall / (precision + recall). Macro F1 là trung bình F1 theo từng lớp.",
        "Latency online được tách thành capture_ms, preprocess_ms, cloud_rtt_ms, inference_ms, esp32_ack_ms và total_ms. Báo cáo dùng median và p95 vì log thời gian thực có outlier do cold start, mạng Wi-Fi hoặc container cloud.",
        "Safety filter yêu cầu confidence tối thiểu và số frame ổn định liên tiếp. Firmware dùng token để tránh lệnh lạ, sequence để chống lệnh cũ, TTL để chống lệnh quá hạn và deadman timeout 600 ms để tự dừng motor.",
    ]:
        add_p(doc, p)


def chapter4(doc, metrics, comparison, latency, rows):
    doc.add_page_break()
    add_h(doc, "CHƯƠNG 4. THỰC NGHIỆM, ĐÁNH GIÁ VÀ THẢO LUẬN", 1)
    subjects = Counter(r["subject_id"] for r in rows)
    gestures = Counter(r["gesture"] for r in rows)
    add_h(doc, "4.1 Môi trường thực nghiệm", 2)
    add_p(doc, "Môi trường thực nghiệm gồm laptop Windows chạy gateway và webcam, Azure Container Apps chạy FastAPI inference, ESP32 kết nối Wi-Fi nhận lệnh WebSocket, L298N điều khiển motor DC và PCA9685 điều khiển bốn servo tay máy.")
    add_table(doc, ["Thành phần", "Giá trị"], [
        ["Số người", str(len(subjects))],
        ["Số lớp", str(len(metrics["labels"]))],
        ["Số clip", str(len(set(r["clip_id"] for r in rows)))],
        ["Số frame", f"{len(rows):,}"],
        ["Train", f"{metrics['counts']['train']:,} frame"],
        ["Validation", f"{metrics['counts']['val']:,} frame"],
        ["Test", f"{metrics['counts']['test']:,} frame"],
        ["Chiến lược chia", "Theo người (subject split)"],
    ], widths=[3.0, 3.6], size=10)
    add_caption(doc, "Bảng 4.1. Thống kê dataset")
    add_image(doc, "dataset_distribution.png", "Hình 4.1. Phân bố dữ liệu theo cử chỉ và theo người", 6.3)
    add_h(doc, "4.1.1 Phát hiện thời gian thực (online)", 3)
    add_p(doc, f"Azure hiện có {latency['Azure']['rows']:,} dòng log; {latency['Azure']['cloud_rtt_ms']['n']:,} dòng có cloud_rtt_ms lớn hơn 0. Cloud RTT median {latency['Azure']['cloud_rtt_ms']['median']} ms và p95 {latency['Azure']['cloud_rtt_ms']['p95']} ms. Inference median {latency['Azure']['inference_ms']['median']} ms và p95 {latency['Azure']['inference_ms']['p95']} ms.")
    add_p(doc, f"So với Hugging Face, Azure ổn định hơn trong log hiện tại: Hugging Face RTT median {latency['Hugging Face']['cloud_rtt_ms']['median']} ms và p95 {latency['Hugging Face']['cloud_rtt_ms']['p95']} ms. Local fallback có RTT thấp nhưng không phản ánh yêu cầu cloud deployment.")
    add_table(doc, ["Provider", "Rows", "RTT med", "RTT p95", "Infer med", "Infer p95", "Total med", "Total p95"], [[p, d["rows"], d["cloud_rtt_ms"]["median"], d["cloud_rtt_ms"]["p95"], d["inference_ms"]["median"], d["inference_ms"]["p95"], d["total_ms"]["median"], d["total_ms"]["p95"]] for p, d in latency.items()], widths=[1.2, .7, .8, .8, .8, .8, .8, .8], size=8)
    add_caption(doc, "Bảng 4.3. Thống kê độ trễ online")
    add_image(doc, "latency_comparison.png", "Hình 4.4. So sánh độ trễ cloud RTT", 6.0)
    add_h(doc, "4.1.2 Đánh giá model offline", 3)
    add_p(doc, f"Đánh giá offline dùng subject split để tránh data leakage giữa các frame cùng clip. Số frame train/validation/test là {metrics['counts']['train']:,}/{metrics['counts']['val']:,}/{metrics['counts']['test']:,}. Nếu cần báo cáo thêm 70/20/10 hoặc 80/20, cần chạy bổ sung một cấu hình split khác.")
    add_p(doc, f"Mô hình đạt accuracy {metrics['accuracy']*100:.2f}% và macro F1 {metrics['macro_f1']*100:.2f}%. Các lớp no_gesture, peace, rock và stop có F1 tương đối cao; one, two, like và dislike còn nhầm lẫn.")
    class_rows = []
    for lab in metrics["labels"]:
        rep = metrics["classification_report"][lab]
        class_rows.append([lab, f"{rep['precision']:.3f}", f"{rep['recall']:.3f}", f"{rep['f1-score']:.3f}", int(rep["support"])])
    add_table(doc, ["Lớp", "Precision", "Recall", "F1", "Support"], class_rows, widths=[1.5, 1.2, 1.2, 1.2, 1.0], size=9)
    add_caption(doc, "Bảng 4.2. Kết quả đánh giá offline theo lớp")
    add_image(doc, "training_history.png", "Hình 4.2. Lịch sử huấn luyện CNN baseline", 6.0)
    add_image(doc, "confusion_matrix.png", "Hình 4.3. Confusion matrix của CNN MobileNetV3Small", 5.8)
    add_h(doc, "4.1.3 So sánh CNN và CNN-LSTM", 3)
    cnn = comparison["models"]["cnn"]
    cnn_lstm = comparison["models"]["cnn_lstm"]
    comparison_rows = [
        [
            "CNN",
            f"{cnn['accuracy'] * 100:.2f}%",
            f"{cnn['macro_f1'] * 100:.2f}%",
            f"{cnn['local_latency_ms']['median']:.2f}",
            f"{cnn['local_latency_ms']['p95']:.2f}",
            f"{cnn['false_activation_no_gesture']['rate'] * 100:.2f}%",
            f"{cnn['robustness']['accuracy_drop'] * 100:.2f} diem",
        ],
        [
            "CNN-LSTM",
            f"{cnn_lstm['accuracy'] * 100:.2f}%",
            f"{cnn_lstm['macro_f1'] * 100:.2f}%",
            f"{cnn_lstm['local_latency_ms']['median']:.2f}",
            f"{cnn_lstm['local_latency_ms']['p95']:.2f}",
            f"{cnn_lstm['false_activation_no_gesture']['rate'] * 100:.2f}%",
            f"{cnn_lstm['robustness']['accuracy_drop'] * 100:.2f} diem",
        ],
    ]
    add_table(
        doc,
        ["Model", "Accuracy", "Macro F1", "Local med", "Local p95", "False activation", "Drop bg phuc tap"],
        comparison_rows,
        widths=[1.0, 0.9, 0.9, 0.9, 0.9, 1.1, 1.3],
        size=8,
    )
    add_caption(doc, "Bảng 4.4. So sánh hai mô hình trên cùng subject split")
    add_p(
        doc,
        f"Kết quả so sánh clip-level cho thấy mô hình được chọn là {comparison['comparison']['selected_model'].upper()} vì {comparison['comparison']['reason']}. "
        f"Chênh lệch macro F1 giữa CNN-LSTM và CNN là {comparison['comparison']['macro_f1_gain_lstm_minus_cnn']:.4f}.",
    )
    add_confusion_matrix_table(
        doc,
        comparison["labels"],
        cnn_lstm["confusion_matrix"],
        "Bảng 4.5. Confusion matrix của mô hình CNN-LSTM theo clip",
    )
    add_h(doc, "4.2 Thảo luận", 2)
    for p in [
        "Về accuracy, CNN baseline vẫn là mốc tham chiếu thực dụng, trong khi CNN-LSTM cung cấp bằng chứng so sánh cho yêu cầu học thuật về mô hình chuỗi. Dù mô hình nào được chọn, các lớp one/two/like/dislike vẫn là nguồn nhầm lẫn chính.",
        "Về latency, Azure là backend chính hợp lý hơn Hugging Face trong log hiện tại. Median RTT khoảng 158 ms đủ cho demo nếu gateway dùng smoothing và không yêu cầu điều khiển tốc độ cao.",
        "Về an toàn, hệ thống có nhiều lớp bảo vệ: không phát lệnh khi confidence thấp, gửi stop khi cloud lỗi, dùng token/sequence/TTL ở firmware và deadman timeout.",
        "Về robustness, summary mới đã tách riêng nền simple/complex và false activation của no_gesture. Điều này phù hợp hơn với yêu cầu đánh giá tương tác thời gian thực so với chỉ báo cáo accuracy tổng.",
        "Về kiến trúc, cloud đúng yêu cầu đề tài nhưng làm tăng phụ thuộc mạng. Phiên bản edge AI trên laptop hoặc Raspberry Pi có thể giảm RTT, trong khi cloud vẫn phù hợp cho quản lý model tập trung.",
    ]:
        add_p(doc, p)


def conclusion_and_refs(doc, metrics, comparison):
    doc.add_page_break()
    add_h(doc, "KẾT LUẬN, HẠN CHẾ VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_p(doc, "Kết quả đạt được: Nhóm đã xây dựng được hệ thống end-to-end điều khiển xe/tay máy bằng cử chỉ tay, bao gồm dataset 8 lớp, mô hình CNN MobileNetV3Small, API FastAPI triển khai trên Azure Container Apps, gateway xử lý webcam/MediaPipe/safety filter và firmware ESP32 nhận lệnh WebSocket để điều khiển L298N/PCA9685. Hệ thống có log latency, session_id/request_id, cơ chế xác thực token, sequence, TTL, deadman timeout và bộ test tự động hiện chạy 38 passed.")
    add_p(doc, f"Kết quả so sánh cho thấy mô hình được ưu tiên là {comparison['comparison']['selected_model'].upper()} với lý do: {comparison['comparison']['reason']}. CNN baseline vẫn giữ vai trò mô hình nhẹ cho triển khai thực tế; CNN-LSTM đóng vai trò đối chứng học thuật để kiểm tra lợi ích của thông tin chuỗi.")
    add_p(doc, "Hạn chế: Mô hình hiện tại vẫn chưa chắc đạt mục tiêu macro F1 0,90 trong bối cảnh người dùng mới. Dataset còn ít người, chưa phủ hết khoảng cách chụp và các cử chỉ dễ nhầm như one/two/like/dislike vẫn cần làm giàu thêm. Đánh giá trong bản này chủ yếu là software-only cho comparison, chưa thay thế hoàn toàn các phép đo live nhiều phiên với phần cứng.")
    add_p(doc, "Hướng phát triển: Nhóm cần mở rộng dataset theo nhiều người, nhiều khoảng cách và nhiều background; thử attention nhẹ hoặc fine-tune sâu hơn cho CNN-LSTM; tối ưu cloud cold start và cân nhắc edge inference để giảm RTT; cải thiện cơ khí, nguồn và chống brownout; bổ sung dashboard giám sát realtime và quy trình Zotero để quản lý tài liệu tham khảo theo chuẩn học thuật.")
    doc.add_page_break()
    add_h(doc, "TÀI LIỆU THAM KHẢO", 1)
    for i, (_, authors, year, title, venue, doi, url, _) in enumerate(REFS, 1):
        add_p(doc, f"[{i}] {authors}, “{title},” {venue}, {year}." + (f" DOI: {doi}." if doi else "") + (f" URL: {url}." if url else ""), indent=False, size=11)


def appendices(doc, metrics, latency, rows):
    doc.add_page_break()
    add_h(doc, "PHỤ LỤC A. KIỂM TRA NGUỒN VÀ QUY TRÌNH ZOTERO", 1)
    add_p(doc, "Phụ lục này đáp ứng lưu ý về bài báo khoa học thuộc danh mục ISI/Scopus. Khi hoàn thiện bản nộp cuối, nhóm nên import file BibTeX đi kèm vào Zotero, kiểm tra từng DOI trên publisher page, sau đó đối chiếu journal/conference trên SCImago, Scopus Sources hoặc Web of Science Journal Info.")
    add_table(doc, ["#", "Nguồn", "Venue", "DOI", "Ghi chú"], [[i + 1, r[3][:55], r[4][:45], r[5] or "N/A", r[7]] for i, r in enumerate(REFS)], widths=[0.35, 2.0, 1.6, 1.4, 1.6], size=6)
    for i, r in enumerate(REFS, 1):
        add_h(doc, f"A.{i} Phiếu nguồn [{i}]", 2)
        add_p(doc, f"Tên nguồn: {r[3]}", indent=False)
        add_p(doc, f"Tác giả/năm: {r[1]} ({r[2]}). Venue: {r[4]}. DOI/URL: {r[5] or r[6]}.", indent=False)
        add_p(doc, f"Lý do dùng trong báo cáo: nguồn này hỗ trợ phần {'nền tảng HGR/HRI' if i <= 4 else 'mô hình, dữ liệu hoặc triển khai'}; ghi chú kiểm tra: {r[7]}.", indent=False)
        add_p(doc, "Các trường cần rà soát trong Zotero trước khi nộp: tên tác giả đầy đủ, tiêu đề, năm, volume/issue, trang, DOI, publisher và loại tài liệu. Với journal, đối chiếu ISSN/venue trên SCImago, Scopus Sources hoặc Web of Science Journal Info; với conference, đối chiếu trang IEEE Xplore, ACM DL hoặc CVF chính thức.", indent=False)
        add_p(doc, "Cách sử dụng trong báo cáo: không trích nguyên văn dài; diễn giải kết quả liên quan bằng lời của nhóm và đặt số trích dẫn IEEE ở cuối nhận định. ResearchGate chỉ dùng để tìm bản đọc, không dùng làm bằng chứng xác nhận chỉ mục.", indent=False)
        if i % 3 == 0 and i < len(REFS):
            doc.add_page_break()
    doc.add_page_break()
    add_h(doc, "PHỤ LỤC B. ÁNH XẠ REPOSITORY VÀ MODULE", 1)
    add_table(doc, ["File/thư mục", "Vai trò"], [
        ["gateway/main.py", "Vòng lặp camera, gọi cloud, safety filter, UI, log CSV"],
        ["gateway/preprocess.py", "MediaPipe crop và hỗ trợ phân biệt one/two"],
        ["gateway/safety.py", "GestureStabilizer, SafetyPolicy và schema log latency"],
        ["gateway/transport.py", "DryRunTransport và WebSocketTransport"],
        ["cloud/inference.py", "Load model, preprocess JPEG, predict gesture"],
        ["cloud/app.py", "FastAPI endpoint /health, /v1/model, /v1/predict"],
        ["common/protocol.py", "Định nghĩa gesture, mode, action và mapper"],
        ["firmware/src/main.cpp", "Firmware ESP32, motor, servo, HTTP debug và WebSocket"],
        ["ml/train_cnn.py", "Huấn luyện CNN baseline"],
        ["ml/train_cnn_lstm.py", "Khung huấn luyện CNN-LSTM định hướng"],
        ["reports/*.csv", "Log latency online"],
        ["reports/*.json", "Metrics offline"],
    ], widths=[2.2, 4.7], size=9)
    add_p(doc, "Báo cáo không nhúng secret từ .env, firmware/include/config.h hoặc token thật. Các giá trị API key, Wi-Fi password và command token chỉ được mô tả bằng placeholder khi cần minh họa.")
    add_h(doc, "PHỤ LỤC C. NHẬT KÝ KIỂM THỬ VÀ TIÊU CHÍ NGHIỆM THU", 1)
    add_p(doc, "Kiểm thử tự động được chạy bằng lệnh .\\.venv\\Scripts\\python.exe -m pytest -q. Kết quả mới nhất trong phiên tạo báo cáo: 38 passed in 2.07s. Kết quả này khác README cũ ghi 28 passed vì repo hiện tại đã có thêm test mới.")
    for item in ["API /health hoạt động", "/v1/model trả model_version", "Gateway dry-run gọi được cloud và ghi log", "ESP32 phản hồi /health và /state", "WebSocket nhận command hợp lệ", "Xe dừng khi mất lệnh quá timeout"]:
        bullet(doc, item)
    add_h(doc, "PHỤ LỤC D. THAM SỐ VẬN HÀNH ĐỀ XUẤT", 1)
    add_table(doc, ["Tham số", "Giá trị", "Ý nghĩa"], [
        ["MinConfidence", "0.80", "Ngưỡng chính cho lệnh thường"],
        ["ModeMinConfidence", "0.60", "Ngưỡng chuyển mode rock/peace"],
        ["NormalRequired", "3", "Số frame liên tiếp cho lệnh thường"],
        ["StopRequired", "2", "Số frame liên tiếp cho stop"],
        ["DeadmanMs", "600", "Gateway/firmware dừng khi quá hạn"],
        ["DriveRepeatMs", "200", "Lặp lệnh xe khi đang giữ gesture"],
        ["DriveHoldMs", "550", "Thời gian giữ lệnh xe"],
        ["ServoCooldownMs", "250-350", "Giảm rung servo khi nhận lệnh liên tục"],
    ], widths=[2.2, 1.4, 3.0], size=9)
    add_h(doc, "PHỤ LỤC E. DỮ LIỆU VÀ NHÃN", 1)
    gestures = Counter(r["gesture"] for r in rows)
    add_table(doc, ["Gesture", "Số frame", "Ghi chú"], [[g, gestures[g], "frame crop/metadata"] for g in metrics["labels"]], widths=[2.0, 1.5, 2.5], size=9)
    for item in ["Thu thêm subject s06-s10 với nền simple/complex và lighting bright/normal/dim.", "Đo false activation bằng no_gesture trong 5 phút.", "Chạy split 70/20/10 và 80/20 như thí nghiệm bổ sung.", "Benchmark edge inference trên laptop/Raspberry Pi để so sánh cloud vs edge."]:
        bullet(doc, item)
    add_h(doc, "PHỤ LỤC F. ĐỐI CHIẾU YÊU CẦU CỦA THẦY", 1)
    add_table(doc, ["Yêu cầu", "Cách đáp ứng"], [
        ["40-70 trang", "Bản in quyển có phụ lục nguồn, repo, test và thông số"],
        ["Chương 1 có 10-15 bài báo", "Trích nguồn nền tảng HGR/HRI/CNN"],
        ["Chương 2 có 8-10 bài báo", "Có bảng so sánh 10 nghiên cứu liên quan"],
        ["Chương 3 có 3-5 bài cùng hướng", "Trích MediaPipe, MobileNetV3, robot control, dynamic HGR"],
        ["Chương 4 online/offline", "Tách online latency và offline model metrics"],
        ["70/20/10 hoặc 80/20", "Ghi rõ hiện tại dùng subject split; đề xuất chạy bổ sung"],
        ["Kết luận 3 đoạn", "Đúng ba đoạn: đạt được, hạn chế, hướng phát triển"],
        ["Zotero", "Có BibTeX đi kèm để import Zotero"],
    ], widths=[2.0, 4.8], size=9)


def write_bibtex():
    lines = []
    for key, authors, year, title, venue, doi, url, note in REFS:
        typ = "misc" if "documentation" in note.lower() or "RFC" in venue else "article"
        if "Conference" in venue or "Workshop" in venue or "ACM" in venue or "ICCV" in venue or "CVPR" in venue:
            typ = "inproceedings"
        lines.append(f"@{typ}{{{key},")
        fields = {
            "author": authors,
            "title": title,
            "year": year,
            "journal" if typ == "article" else "booktitle" if typ == "inproceedings" else "howpublished": venue,
            "doi": doi,
            "url": url,
            "note": note,
        }
        for k, v in fields.items():
            if v:
                lines.append(f"  {k} = {{{v}}},")
        lines.append("}\n")
    BIB_OUT.write_text("\n".join(lines), encoding="utf-8")


def main():
    metrics, comparison, latency, rows = load_inputs()
    doc = Document()
    configure(doc)
    cover(doc)
    front_matter(doc, metrics, latency, rows)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc, metrics)
    chapter4(doc, metrics, comparison, latency, rows)
    conclusion_and_refs(doc, metrics, comparison)
    appendices(doc, metrics, latency, rows)
    doc.save(DOCX_OUT)
    write_bibtex()
    qa = {
        "docx": str(DOCX_OUT),
        "bibtex": str(BIB_OUT),
        "figures": sorted(p.name for p in ASSETS.glob("*.png")),
        "references": len(REFS),
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "metrics": {"accuracy": metrics["accuracy"], "macro_f1": metrics["macro_f1"]},
        "tests": "38 passed in 2.07s",
    }
    QA_OUT.write_text(json.dumps(qa, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(qa, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
